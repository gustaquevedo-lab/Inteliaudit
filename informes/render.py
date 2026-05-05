"""
Generación de informes de auditoría en Word y PDF.
Usa Jinja2 para templates HTML → WeasyPrint para PDF, python-docx para Word.
"""
from pathlib import Path
from typing import Optional

from jinja2 import Environment, FileSystemLoader, select_autoescape
from rich.console import Console
from sqlalchemy.ext.asyncio import AsyncSession

from analisis.riesgo import formatear_pyg, resumir_contingencias
from config.settings import settings
from db import db as crud

console = Console()

TEMPLATES_DIR = Path(__file__).parent / "templates"


class RenderInforme:
    """
    Genera los diferentes tipos de informe de auditoría.
    """

    def __init__(self, db: AsyncSession):
        self.db = db
        self.jinja = Environment(
            loader=FileSystemLoader(str(TEMPLATES_DIR)),
            autoescape=select_autoescape(["html"]),
        )
        self.jinja.filters["pyg"] = formatear_pyg

    async def generar_informe_auditoria(
        self,
        auditoria_id: str,
        output_dir: Optional[Path] = None,
    ) -> dict[str, Path]:
        """
        Genera el informe completo de auditoría impositiva en Word y PDF.

        Returns:
            Dict con keys 'docx' y 'pdf' apuntando a los archivos generados.
        """
        auditoria = await crud.get_auditoria(self.db, auditoria_id)
        if not auditoria:
            raise ValueError(f"Auditoría {auditoria_id} no encontrada")

        cliente = await crud.get_cliente(self.db, auditoria.cliente_ruc)
        hallazgos = await crud.get_hallazgos(self.db, auditoria_id)
        hallazgos_data = _serializar_hallazgos(hallazgos)
        resumen = resumir_contingencias(hallazgos_data)

        import json
        contexto = {
            "auditoria": {
                "id": auditoria.id,
                "periodo_desde": auditoria.periodo_desde,
                "periodo_hasta": auditoria.periodo_hasta,
                "impuestos": json.loads(auditoria.impuestos),
                "auditor": auditoria.auditor or "Inteliaudit",
                "fecha_inicio": auditoria.fecha_inicio,
            },
            "cliente": {
                "ruc": cliente.ruc,
                "razon_social": cliente.razon_social,
                "actividad": cliente.actividad_principal,
                "regimen": cliente.regimen,
                "direccion": cliente.direccion,
            },
            "hallazgos": hallazgos_data,
            "resumen": resumen,
            "fecha_informe": __import__("datetime").date.today().isoformat(),
        }

        output = output_dir or Path(settings.storage_path) / cliente.ruc / "informes"
        output.mkdir(parents=True, exist_ok=True)

        base_nombre = f"informe_auditoria_{auditoria_id[:8]}_{auditoria.periodo_desde}_{auditoria.periodo_hasta}"

        paths: dict[str, Path] = {}
        paths["html"] = await self._renderizar_html(contexto, "informe_auditoria.html", output / f"{base_nombre}.html")
        paths["pdf"] = await self._html_a_pdf(paths["html"], output / f"{base_nombre}.pdf")
        paths["docx"] = await self._generar_docx(contexto, output / f"{base_nombre}.docx")

        # Registrar en DB
        from db.models import Informe
        informe = Informe(
            auditoria_id=auditoria_id,
            tipo="auditoria_impositiva",
            archivo_docx=str(paths["docx"]),
            archivo_pdf=str(paths["pdf"]),
        )
        self.db.add(informe)
        await self.db.flush()

        console.print(f"[green]✓[/] Informe generado: {paths['pdf'].name}")
        return paths

    # --------------------------------------------------------
    #  Métodos internos
    # --------------------------------------------------------

    async def _renderizar_html(self, contexto: dict, template_name: str, destino: Path) -> Path:
        """Renderiza un template Jinja2 a HTML."""
        try:
            template = self.jinja.get_template(template_name)
        except Exception:
            # Template no existe aún — usar template básico inline
            html = _template_basico(contexto)
            destino.write_text(html, encoding="utf-8")
            return destino

        html = template.render(**contexto)
        destino.write_text(html, encoding="utf-8")
        return destino

    async def _html_a_pdf(self, html_path: Path, destino: Path) -> Path:
        """Convierte HTML a PDF con WeasyPrint."""
        try:
            from weasyprint import HTML
            HTML(filename=str(html_path)).write_pdf(str(destino))
        except ImportError:
            console.print("[yellow]⚠[/] WeasyPrint no disponible. PDF no generado.")
            destino = html_path
        return destino

    async def _generar_docx(self, contexto: dict, destino: Path) -> Path:
        """Genera informe Word con python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            doc.add_heading("INFORME DE AUDITORÍA IMPOSITIVA", level=0)
            doc.add_paragraph(f"Cliente: {contexto['cliente']['razon_social']}")
            doc.add_paragraph(f"RUC: {contexto['cliente']['ruc']}")
            doc.add_paragraph(f"Período: {contexto['auditoria']['periodo_desde']} a {contexto['auditoria']['periodo_hasta']}")
            doc.add_paragraph(f"Fecha: {contexto['fecha_informe']}")

            doc.add_heading("Resumen de Contingencias", level=1)
            resumen = contexto["resumen"]
            doc.add_paragraph(f"Total contingencia estimada: Gs. {formatear_pyg(resumen['total_contingencia'])}")
            doc.add_paragraph(f"Impuesto omitido: Gs. {formatear_pyg(resumen['total_impuesto'])}")
            doc.add_paragraph(f"Multas estimadas: Gs. {formatear_pyg(resumen['total_multa'])}")
            doc.add_paragraph(f"Intereses estimados: Gs. {formatear_pyg(resumen['total_intereses'])}")

            doc.add_heading("Hallazgos", level=1)
            for h in contexto["hallazgos"]:
                if h.get("estado") == "descartado":
                    continue
                doc.add_heading(f"{h['tipo_hallazgo']} — {h['periodo']}", level=2)
                doc.add_paragraph(h["descripcion"])
                doc.add_paragraph(f"Contingencia: Gs. {formatear_pyg(h['total_contingencia'])} | Riesgo: {h['nivel_riesgo'].upper()}")
                doc.add_paragraph(f"Base legal: {h['articulo_legal']}")

            doc.save(str(destino))
        except ImportError:
            console.print("[yellow]⚠[/] python-docx no disponible. DOCX no generado.")

        return destino


# --------------------------------------------------------
#  Helpers
# --------------------------------------------------------

def _serializar_hallazgos(hallazgos) -> list[dict]:
    import json
    return [
        {
            "id": h.id,
            "impuesto": h.impuesto,
            "periodo": h.periodo,
            "tipo_hallazgo": h.tipo_hallazgo,
            "descripcion": h.descripcion,
            "articulo_legal": h.articulo_legal,
            "impuesto_omitido": h.impuesto_omitido,
            "multa_estimada": h.multa_estimada,
            "intereses_estimados": h.intereses_estimados,
            "total_contingencia": h.total_contingencia,
            "nivel_riesgo": h.nivel_riesgo,
            "estado": h.estado,
        }
        for h in hallazgos
    ]


def _template_basico(ctx: dict) -> str:
    """Template HTML mínimo mientras no existan los archivos Jinja2."""
    r = ctx["resumen"]
    hallazgos_html = "".join(
        f"<tr><td>{h['impuesto']}</td><td>{h['periodo']}</td><td>{h['tipo_hallazgo']}</td>"
        f"<td style='text-align:right'>{formatear_pyg(h['total_contingencia'])}</td>"
        f"<td>{h['nivel_riesgo'].upper()}</td></tr>"
        for h in ctx["hallazgos"]
        if h.get("estado") != "descartado"
    )
    return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>Informe de Auditoría — {ctx['cliente']['razon_social']}</title>
<style>
  body {{ font-family: Helvetica Neue, Arial, sans-serif; color: #091624; margin: 40px; }}
  h1 {{ color: #2E84F0; }}
  h2 {{ color: #1558B0; border-bottom: 2px solid #2E84F0; padding-bottom: 4px; }}
  table {{ border-collapse: collapse; width: 100%; margin-top: 16px; }}
  th {{ background: #2E84F0; color: white; padding: 8px 12px; text-align: left; }}
  td {{ padding: 8px 12px; border-bottom: 1px solid #E4EAF4; }}
  .total {{ font-size: 1.2em; font-weight: bold; color: #E53E3E; }}
</style>
</head>
<body>
<h1>Informe de Auditoría Impositiva</h1>
<p><strong>Cliente:</strong> {ctx['cliente']['razon_social']} — RUC {ctx['cliente']['ruc']}</p>
<p><strong>Período:</strong> {ctx['auditoria']['periodo_desde']} a {ctx['auditoria']['periodo_hasta']}</p>
<p><strong>Fecha:</strong> {ctx['fecha_informe']}</p>

<h2>Resumen de Contingencias</h2>
<p class="total">Total contingencia estimada: Gs. {formatear_pyg(r['total_contingencia'])}</p>
<p>Impuesto omitido: Gs. {formatear_pyg(r['total_impuesto'])} |
   Multas: Gs. {formatear_pyg(r['total_multa'])} |
   Intereses: Gs. {formatear_pyg(r['total_intereses'])}</p>

<h2>Hallazgos Identificados</h2>
<table>
  <thead><tr><th>Impuesto</th><th>Período</th><th>Tipo</th><th>Contingencia (Gs.)</th><th>Riesgo</th></tr></thead>
  <tbody>{hallazgos_html}</tbody>
</table>
</body>
</html>"""
