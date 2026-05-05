"""
Generador de informe Word profesional — Inteliaudit.
Estructura legal completa para auditoría impositiva paraguaya.
python-docx con estilos de marca Inteliaudit.
"""
import json
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

from analisis.riesgo import formatear_pyg, resumir_contingencias

# ============================================================
#  Paleta Inteliaudit
# ============================================================
AZUL       = RGBColor(0x2E, 0x84, 0xF0)
AZUL_OSCURO = RGBColor(0x15, 0x58, 0xB0)
VERDE      = RGBColor(0x22, 0xC4, 0x7E)
NAVY       = RGBColor(0x09, 0x16, 0x24)
GRIS       = RGBColor(0xA8, 0xB4, 0xC8)
ROJO       = RGBColor(0xE5, 0x3E, 0x3E)
AMBAR      = RGBColor(0xD9, 0x77, 0x06)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)
BORDE      = RGBColor(0xE4, 0xEA, 0xF4)
FONDO      = RGBColor(0xF5, 0xF7, 0xFB)

COLORES_RIESGO = {"alto": ROJO, "medio": AMBAR, "bajo": VERDE}


def generar_informe_word(
    auditoria: dict,
    cliente: dict,
    hallazgos: list[dict],
    logo_cliente_path: Optional[Path] = None,
    logo_inteliaudit_path: Optional[Path] = None,
    notas_auditor: Optional[str] = None,
) -> bytes:
    """
    Genera el informe de auditoría impositiva completo en formato Word.

    Args:
        auditoria: Dict con id, periodo_desde, periodo_hasta, impuestos, auditor, materialidad
        cliente: Dict con ruc, razon_social, actividad_principal, regimen, direccion
        hallazgos: Lista de dicts con campos de la tabla hallazgos
        logo_cliente_path: Path al logo del cliente (PNG/JPG) — opcional
        logo_inteliaudit_path: Path al logo Inteliaudit — opcional
        notas_auditor: Observaciones generales adicionales del auditor

    Returns:
        Bytes del archivo .docx
    """
    doc = Document()
    _configurar_documento(doc)
    _aplicar_estilos(doc)

    hallazgos_activos = [h for h in hallazgos if h.get("estado") != "descartado"]
    resumen = resumir_contingencias(hallazgos_activos)

    # Estructura del informe
    _portada(doc, auditoria, cliente, logo_cliente_path, logo_inteliaudit_path)
    _salto_pagina(doc)
    _indice(doc)
    _salto_pagina(doc)
    _seccion_resumen_ejecutivo(doc, auditoria, cliente, hallazgos_activos, resumen)
    _salto_pagina(doc)
    _seccion_datos_contribuyente(doc, cliente, auditoria)
    _salto_pagina(doc)
    _seccion_metodologia(doc, auditoria)
    _salto_pagina(doc)
    _seccion_hallazgos(doc, hallazgos_activos, auditoria)
    _salto_pagina(doc)
    _seccion_matriz_riesgo(doc, hallazgos_activos)
    _salto_pagina(doc)
    _seccion_conclusiones(doc, resumen, auditoria, cliente, notas_auditor)
    _pie_pagina(doc, auditoria)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
#  Configuración general del documento
# ============================================================

def _configurar_documento(doc: Document) -> None:
    secciones = doc.sections
    for seccion in secciones:
        seccion.top_margin = Cm(2.5)
        seccion.bottom_margin = Cm(2.5)
        seccion.left_margin = Cm(3.0)
        seccion.right_margin = Cm(2.5)
        seccion.page_width = Cm(21.59)
        seccion.page_height = Cm(27.94)


def _aplicar_estilos(doc: Document) -> None:
    """Define estilos personalizados reutilizables."""
    estilos = doc.styles

    # Estilo Normal base
    normal = estilos["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(6)

    # Título 1 — secciones principales
    for nivel, size, color in [("Heading 1", 16, AZUL), ("Heading 2", 13, AZUL_OSCURO), ("Heading 3", 11, NAVY)]:
        try:
            h = estilos[nivel]
            h.font.name = "Arial"
            h.font.bold = True
            h.font.size = Pt(size)
            h.font.color.rgb = color
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        except Exception:
            pass


# ============================================================
#  Portada
# ============================================================

def _portada(doc, auditoria, cliente, logo_cliente, logo_inteliaudit):
    # Banda superior con color de marca
    p = doc.add_paragraph()
    _set_paragraph_background(p, AZUL)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(" " * 200)
    run.font.size = Pt(6)

    doc.add_paragraph()

    # Logo Inteliaudit (si existe)
    if logo_inteliaudit and logo_inteliaudit.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_logo.add_run()
        run.add_picture(str(logo_inteliaudit), width=Cm(4))
    else:
        p_brand = doc.add_paragraph()
        p_brand.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p_brand.add_run("Inteli")
        r1.font.bold = True
        r1.font.size = Pt(20)
        r1.font.color.rgb = AZUL
        r2 = p_brand.add_run("audit")
        r2.font.bold = False
        r2.font.size = Pt(20)
        r2.font.color.rgb = VERDE

    doc.add_paragraph()
    doc.add_paragraph()

    # Logo del cliente
    if logo_cliente and logo_cliente.exists():
        p_logo_cl = doc.add_paragraph()
        p_logo_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo_cl.add_run()
        run.add_picture(str(logo_cliente), height=Cm(3))
        doc.add_paragraph()

    # Título principal
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("INFORME DE AUDITORÍA IMPOSITIVA")
    r.font.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = AZUL

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitulo.add_run("CONFIDENCIAL")
    r2.font.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = ROJO

    doc.add_paragraph()

    # Tabla de datos del encargo
    tabla = doc.add_table(rows=5, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    impuestos = auditoria.get("impuestos", [])
    if isinstance(impuestos, str):
        impuestos = json.loads(impuestos)

    datos = [
        ("Cliente", cliente.get("razon_social", "")),
        ("RUC", cliente.get("ruc", "")),
        ("Período auditado", f"{auditoria.get('periodo_desde', '')} a {auditoria.get('periodo_hasta', '')}"),
        ("Impuestos auditados", ", ".join(impuestos)),
        ("Fecha del informe", date.today().strftime("%d de %B de %Y")),
    ]

    for i, (label, valor) in enumerate(datos):
        fila = tabla.rows[i]
        c0 = fila.cells[0]
        c1 = fila.cells[1]
        _set_cell_background(c0, AZUL)
        r_label = c0.paragraphs[0].add_run(label)
        r_label.font.bold = True
        r_label.font.color.rgb = BLANCO
        r_label.font.size = Pt(10)
        r_valor = c1.paragraphs[0].add_run(valor)
        r_valor.font.size = Pt(10)
        r_valor.font.color.rgb = NAVY

    # Auditor
    if auditoria.get("auditor"):
        doc.add_paragraph()
        p_aud = doc.add_paragraph()
        p_aud.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_aud = p_aud.add_run(f"Preparado por: {auditoria['auditor']}")
        r_aud.font.italic = True
        r_aud.font.size = Pt(9)
        r_aud.font.color.rgb = GRIS

    # Línea inferior de marca
    doc.add_paragraph()
    doc.add_paragraph()
    p_bottom = doc.add_paragraph()
    _set_paragraph_background(p_bottom, VERDE)
    run = p_bottom.add_run(" " * 200)
    run.font.size = Pt(4)


# ============================================================
#  Índice
# ============================================================

def _indice(doc):
    doc.add_heading("Índice de Contenidos", level=1)
    secciones = [
        "1. Resumen Ejecutivo",
        "2. Datos del Contribuyente",
        "3. Metodología de Auditoría",
        "4. Hallazgos por Impuesto",
        "5. Matriz de Riesgo",
        "6. Conclusiones y Recomendaciones",
    ]
    for s in secciones:
        p = doc.add_paragraph(s, style="List Number")
        p.paragraph_format.space_after = Pt(4)


# ============================================================
#  1. Resumen Ejecutivo
# ============================================================

def _seccion_resumen_ejecutivo(doc, auditoria, cliente, hallazgos, resumen):
    doc.add_heading("1. Resumen Ejecutivo", level=1)

    impuestos = auditoria.get("impuestos", [])
    if isinstance(impuestos, str):
        impuestos = json.loads(impuestos)

    p = doc.add_paragraph()
    p.add_run(
        f"En cumplimiento del encargo de auditoría impositiva encomendado, se procedió al análisis "
        f"de las obligaciones tributarias de {cliente.get('razon_social')} "
        f"(RUC: {cliente.get('ruc')}), correspondiente al período "
        f"{auditoria.get('periodo_desde')} a {auditoria.get('periodo_hasta')}, "
        f"abarcando los impuestos: {', '.join(impuestos)}. "
        f"La auditoría fue realizada conforme a las normas de auditoría impositiva vigentes "
        f"en la República del Paraguay."
    )

    # Tabla resumen de contingencias
    doc.add_heading("Resumen de Contingencias", level=2)

    if not resumen["por_impuesto"]:
        doc.add_paragraph("No se identificaron contingencias tributarias significativas en el período auditado.")
        return

    tabla = doc.add_table(rows=1, cols=6)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    encabezados = ["Impuesto", "Hallazgos", "Impuesto omitido (Gs.)", "Multas (Gs.)", "Intereses (Gs.)", "Total contingencia (Gs.)"]
    fila_enc = tabla.rows[0]
    for i, texto in enumerate(encabezados):
        cell = fila_enc.cells[i]
        _set_cell_background(cell, AZUL)
        r = cell.paragraphs[0].add_run(texto)
        r.font.bold = True
        r.font.color.rgb = BLANCO
        r.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for imp, datos in resumen["por_impuesto"].items():
        fila = tabla.add_row()
        valores = [
            imp,
            str(datos["cantidad"]),
            formatear_pyg(datos["impuesto_omitido"]),
            formatear_pyg(datos["multa"]),
            formatear_pyg(datos["intereses"]),
            formatear_pyg(datos["total"]),
        ]
        for j, val in enumerate(valores):
            cell = fila.cells[j]
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            if j == 5:
                r.font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # Fila de totales
    fila_total = tabla.add_row()
    totales = [
        "TOTAL",
        str(resumen["cantidad_hallazgos"]),
        formatear_pyg(resumen["total_impuesto"]),
        formatear_pyg(resumen["total_multa"]),
        formatear_pyg(resumen["total_intereses"]),
        formatear_pyg(resumen["total_contingencia"]),
    ]
    for j, val in enumerate(totales):
        cell = fila_total.cells[j]
        _set_cell_background(cell, NAVY)
        r = cell.paragraphs[0].add_run(val)
        r.font.bold = True
        r.font.color.rgb = BLANCO
        r.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    p_total = doc.add_paragraph()
    r_t = p_total.add_run(f"Contingencia tributaria total estimada: Gs. {formatear_pyg(resumen['total_contingencia'])}")
    r_t.font.bold = True
    r_t.font.size = Pt(11)
    r_t.font.color.rgb = ROJO if resumen["total_contingencia"] > 0 else VERDE


# ============================================================
#  2. Datos del Contribuyente
# ============================================================

def _seccion_datos_contribuyente(doc, cliente, auditoria):
    doc.add_heading("2. Datos del Contribuyente", level=1)

    tabla = doc.add_table(rows=6, cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT

    campos = [
        ("RUC", cliente.get("ruc", "")),
        ("Razón Social", cliente.get("razon_social", "")),
        ("Nombre de Fantasía", cliente.get("nombre_fantasia") or "—"),
        ("Actividad Principal", cliente.get("actividad_principal") or "—"),
        ("Régimen Tributario", (cliente.get("regimen") or "").upper()),
        ("Domicilio Fiscal", cliente.get("direccion") or "—"),
    ]

    for i, (label, valor) in enumerate(campos):
        fila = tabla.rows[i]
        _set_cell_background(fila.cells[0], FONDO)
        r_l = fila.cells[0].paragraphs[0].add_run(label)
        r_l.font.bold = True
        r_l.font.size = Pt(10)
        r_l.font.color.rgb = AZUL_OSCURO
        fila.cells[0].width = Cm(5)
        r_v = fila.cells[1].paragraphs[0].add_run(valor)
        r_v.font.size = Pt(10)
        r_v.font.color.rgb = NAVY


# ============================================================
#  3. Metodología
# ============================================================

def _seccion_metodologia(doc, auditoria):
    doc.add_heading("3. Metodología de Auditoría", level=1)

    doc.add_heading("3.1 Alcance", level=2)
    impuestos = auditoria.get("impuestos", [])
    if isinstance(impuestos, str):
        impuestos = json.loads(impuestos)
    materialidad = auditoria.get("materialidad", 0)

    p = doc.add_paragraph(
        f"La auditoría abarcó los impuestos {', '.join(impuestos)} correspondientes al período "
        f"{auditoria.get('periodo_desde')} a {auditoria.get('periodo_hasta')}. "
        f"Se aplicó un umbral de materialidad de Gs. {formatear_pyg(materialidad)} para la selección de hallazgos a reportar."
    )

    doc.add_heading("3.2 Fuentes de información consultadas", level=2)
    fuentes = [
        ("Portal Marangatú (DNIT)", "Declaraciones juradas presentadas, RG 90, HECHAUKA, estado de cuenta."),
        ("SIFEN (e-Kuatia)", "Validación de comprobantes electrónicos mediante CDC."),
        ("HECHAUKA", "Información declarada por terceros sobre operaciones con el contribuyente."),
        ("Documentación aportada por el contribuyente", "Estados contables, movimientos bancarios, comprobantes adicionales."),
    ]
    for fuente, desc in fuentes:
        p = doc.add_paragraph(style="List Bullet")
        r_f = p.add_run(f"{fuente}: ")
        r_f.font.bold = True
        p.add_run(desc)

    doc.add_heading("3.3 Procedimientos aplicados", level=2)
    procedimientos = [
        "Cruce de totales RG 90 vs. Formulario 120 (crédito fiscal y débito fiscal).",
        "Validación de CDCs de comprobantes electrónicos en SIFEN.",
        "Cruce de comprobantes electrónicos recibidos en SIFEN vs. RG 90 compras.",
        "Cruce de ventas declaradas en RG 90 vs. HECHAUKA recibido.",
        "Verificación del estado de RUC de proveedores.",
        "Análisis de deducibilidad de gastos e IRE.",
        "Verificación de retenciones practicadas vs. HECHAUKA.",
    ]
    for proc in procedimientos:
        doc.add_paragraph(proc, style="List Bullet")

    doc.add_heading("3.4 Marco legal aplicable", level=2)
    normas = [
        "Ley N° 6380/2019 — Modernización y Simplificación del Sistema Tributario Nacional",
        "Ley N° 125/1991 — Código Tributario y sus modificaciones",
        "Decreto N° 3107/2019 — Reglamentación del IRE",
        "Decreto N° 3181/2019 — Reglamentación del IVA bajo Ley 6380",
        "Resolución General N° 69/2020 — Implementación Factura Electrónica (e-Kuatia)",
        "Resolución General N° 80/2021 — Obligatoriedad CDC en comprobantes electrónicos",
        "Resolución General N° 90/2021 — Detalle de comprobantes en Declaración IVA",
    ]
    for norma in normas:
        doc.add_paragraph(norma, style="List Bullet")


# ============================================================
#  4. Hallazgos
# ============================================================

def _seccion_hallazgos(doc, hallazgos, auditoria):
    doc.add_heading("4. Hallazgos por Impuesto", level=1)

    if not hallazgos:
        doc.add_paragraph("No se identificaron hallazgos significativos en el período auditado.")
        return

    # Agrupar por impuesto
    por_impuesto: dict = {}
    for h in hallazgos:
        imp = h.get("impuesto", "OTRO")
        por_impuesto.setdefault(imp, []).append(h)

    seccion_num = 1
    for impuesto, items in por_impuesto.items():
        doc.add_heading(f"4.{seccion_num} {impuesto}", level=2)
        seccion_num += 1

        for idx, h in enumerate(items, 1):
            # Encabezado del hallazgo con color de riesgo
            p_h = doc.add_paragraph()
            riesgo = h.get("nivel_riesgo", "medio")
            color_riesgo = COLORES_RIESGO.get(riesgo, AMBAR)

            r_num = p_h.add_run(f"Hallazgo {idx}: ")
            r_num.font.bold = True
            r_num.font.color.rgb = color_riesgo
            r_num.font.size = Pt(11)

            r_tipo = p_h.add_run(h.get("tipo_hallazgo", "").replace("_", " "))
            r_tipo.font.bold = True
            r_tipo.font.size = Pt(11)
            r_tipo.font.color.rgb = NAVY

            # Tabla de datos del hallazgo
            tabla = doc.add_table(rows=8, cols=2)
            tabla.style = "Table Grid"

            campos_h = [
                ("Período", h.get("periodo", "")),
                ("Nivel de riesgo", h.get("nivel_riesgo", "").upper()),
                ("Base del ajuste", f"Gs. {formatear_pyg(h.get('base_ajuste', 0))}"),
                ("Impuesto omitido", f"Gs. {formatear_pyg(h.get('impuesto_omitido', 0))}"),
                ("Multa estimada (50%)", f"Gs. {formatear_pyg(h.get('multa_estimada', 0))}"),
                ("Intereses estimados (1%/mes)", f"Gs. {formatear_pyg(h.get('intereses_estimados', 0))}"),
                ("CONTINGENCIA TOTAL", f"Gs. {formatear_pyg(h.get('total_contingencia', 0))}"),
                ("Base legal", h.get("articulo_legal", "")),
            ]

            for i, (label, valor) in enumerate(campos_h):
                fila = tabla.rows[i]
                _set_cell_background(fila.cells[0], FONDO)
                r_l = fila.cells[0].paragraphs[0].add_run(label)
                r_l.font.bold = True
                r_l.font.size = Pt(9)
                r_l.font.color.rgb = AZUL_OSCURO
                fila.cells[0].width = Cm(5.5)
                r_v = fila.cells[1].paragraphs[0].add_run(valor)
                r_v.font.size = Pt(9)

                if label == "CONTINGENCIA TOTAL":
                    _set_cell_background(fila.cells[0], NAVY)
                    _set_cell_background(fila.cells[1], NAVY)
                    r_l.font.color.rgb = BLANCO
                    r_v.font.bold = True
                    r_v.font.color.rgb = BLANCO
                elif label == "Nivel de riesgo":
                    r_v.font.color.rgb = color_riesgo
                    r_v.font.bold = True

            # Descripción
            doc.add_heading("Descripción", level=3)
            doc.add_paragraph(h.get("descripcion", ""))

            if h.get("descripcion_tecnica"):
                doc.add_heading("Detalle técnico", level=3)
                doc.add_paragraph(h.get("descripcion_tecnica"))

            if h.get("notas_auditor"):
                doc.add_heading("Observaciones del auditor", level=3)
                p_nota = doc.add_paragraph(h.get("notas_auditor"))
                p_nota.runs[0].font.italic = True

            doc.add_paragraph()


# ============================================================
#  5. Matriz de Riesgo
# ============================================================

def _seccion_matriz_riesgo(doc, hallazgos):
    doc.add_heading("5. Matriz de Riesgo", level=1)

    if not hallazgos:
        doc.add_paragraph("No se identificaron hallazgos a incluir en la matriz de riesgo.")
        return

    tabla = doc.add_table(rows=1, cols=7)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    encabezados = ["#", "Impuesto", "Período", "Tipo", "Contingencia (Gs.)", "Riesgo", "Estado"]
    fila_enc = tabla.rows[0]
    for i, texto in enumerate(encabezados):
        cell = fila_enc.cells[i]
        _set_cell_background(cell, NAVY)
        r = cell.paragraphs[0].add_run(texto)
        r.font.bold = True
        r.font.color.rgb = BLANCO
        r.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    hallazgos_ord = sorted(hallazgos, key=lambda h: h.get("total_contingencia", 0), reverse=True)

    for idx, h in enumerate(hallazgos_ord, 1):
        fila = tabla.add_row()
        riesgo = h.get("nivel_riesgo", "medio")
        color = COLORES_RIESGO.get(riesgo, AMBAR)

        valores = [
            str(idx),
            h.get("impuesto", ""),
            h.get("periodo", ""),
            h.get("tipo_hallazgo", "").replace("_", " "),
            formatear_pyg(h.get("total_contingencia", 0)),
            h.get("nivel_riesgo", "").upper(),
            h.get("estado", "").upper(),
        ]

        for j, val in enumerate(valores):
            cell = fila.cells[j]
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 3 else WD_ALIGN_PARAGRAPH.LEFT

            if j == 5:
                r.font.bold = True
                r.font.color.rgb = color

            if idx % 2 == 0:
                _set_cell_background(cell, FONDO)


# ============================================================
#  6. Conclusiones
# ============================================================

def _seccion_conclusiones(doc, resumen, auditoria, cliente, notas_auditor):
    doc.add_heading("6. Conclusiones y Recomendaciones", level=1)

    doc.add_heading("6.1 Conclusión general", level=2)
    total = resumen.get("total_contingencia", 0)
    cantidad = resumen.get("cantidad_hallazgos", 0)
    alto = resumen.get("por_riesgo", {}).get("alto", 0)

    if total == 0:
        conclusion = (
            f"Como resultado de la auditoría impositiva realizada a {cliente.get('razon_social')} "
            f"por el período {auditoria.get('periodo_desde')} a {auditoria.get('periodo_hasta')}, "
            f"no se identificaron contingencias tributarias significativas. "
            f"Las declaraciones juradas presentadas ante la SET se encuentran en concordancia "
            f"con los datos obtenidos de Marangatú, SIFEN y HECHAUKA."
        )
    else:
        nivel_general = "ALTO" if alto > 0 else ("MEDIO" if resumen.get("por_riesgo", {}).get("medio", 0) > 0 else "BAJO")
        conclusion = (
            f"Como resultado de la auditoría impositiva realizada a {cliente.get('razon_social')} "
            f"(RUC: {cliente.get('ruc')}) por el período {auditoria.get('periodo_desde')} a "
            f"{auditoria.get('periodo_hasta')}, se identificaron {cantidad} hallazgo(s) con una "
            f"contingencia tributaria total estimada de Gs. {formatear_pyg(total)}, "
            f"correspondiente a impuesto omitido, multas e intereses calculados conforme al "
            f"Art. 175 de la Ley N° 125/1991. El nivel de riesgo global de la auditoría es {nivel_general}."
        )

    doc.add_paragraph(conclusion)

    doc.add_heading("6.2 Recomendaciones", level=2)
    recomendaciones = [
        "Regularizar de forma voluntaria los ajustes identificados ante la SET, lo que puede reducir las multas aplicables según Art. 176 Ley 125/1991.",
        "Implementar controles internos para verificar el estado activo de RUCs de proveedores antes de imputar crédito fiscal.",
        "Asegurar que todos los comprobantes de compra a partir de enero 2022 cuenten con CDC válido en SIFEN.",
        "Revisar mensualmente la consistencia entre los totales de RG 90 y el Formulario 120 antes de la presentación.",
        "Conservar la documentación de respaldo por el plazo de prescripción de 5 años conforme a la Ley 125/1991.",
    ]
    for rec in recomendaciones:
        doc.add_paragraph(rec, style="List Bullet")

    if notas_auditor:
        doc.add_heading("6.3 Observaciones adicionales del auditor", level=2)
        doc.add_paragraph(notas_auditor)

    # Firma
    doc.add_paragraph()
    doc.add_paragraph()
    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_firma.add_run(f"Asunción, {date.today().strftime('%d de %B de %Y')}")
    r.font.italic = True
    r.font.size = Pt(10)

    if auditoria.get("auditor"):
        doc.add_paragraph()
        p_aud = doc.add_paragraph()
        p_aud.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_a = p_aud.add_run("_" * 30)
        r_a.font.color.rgb = NAVY
        p_nombre = doc.add_paragraph()
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_n = p_nombre.add_run(auditoria["auditor"])
        r_n.font.bold = True
        r_n.font.size = Pt(10)
        p_cargo = doc.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_c = p_cargo.add_run("Auditor Impositivo")
        r_c.font.size = Pt(9)
        r_c.font.color.rgb = GRIS


# ============================================================
#  Pie de página
# ============================================================

def _pie_pagina(doc, auditoria):
    seccion = doc.sections[0]
    footer = seccion.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Inteliaudit — Informe Confidencial | "
        f"{auditoria.get('periodo_desde')} a {auditoria.get('periodo_hasta')} | "
        f"Generado el {date.today().isoformat()}"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = GRIS


# ============================================================
#  Helpers de formato
# ============================================================

def _salto_pagina(doc):
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _set_paragraph_background(paragraph, color: RGBColor):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_cell_background(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
